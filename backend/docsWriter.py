import os
import ObjectIdentifier
from docx import Document

TRANSLATED_FOLDER = 'translated'

def existing_docs_purger():
    try:
        # List all files in the folder
        for filename in os.listdir(TRANSLATED_FOLDER):
            file_path = os.path.join(TRANSLATED_FOLDER, filename)
            # Check if it's a file and not a directory
            if os.path.isfile(file_path):
                os.remove(file_path)  # Remove the file
        print(f"All files removed from {TRANSLATED_FOLDER}")
    except Exception as e:
        print(f"Error: {e}")

def writeDocs(description, language_table):
    existing_docs_purger()

    translated_desc_dict = ObjectIdentifier.languageWriter(description=description, language_dict=language_table)

    for lang, translation in translated_desc_dict.items():
        # Create a new Document
        doc = Document()

        # Add a heading (level 1)
        doc.add_heading(f'{lang} translation of description', level=1)

        # Add the translated description as a paragraph
        doc.add_paragraph(translation)

        # Define the file path
        file_path = os.path.join(TRANSLATED_FOLDER, f'{lang}_translation.docx')

        # Save the document
        doc.save(file_path)

        print(f'Document saved: {file_path}')

    print(translated_desc_dict)
